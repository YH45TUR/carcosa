"""
Sistema Legal CO - DOCX Parser
Extracción de texto y metadatos de documentos Word usando python-docx.
"""
from typing import Optional, List, Dict, Any
import os


class DOCXParser:
    """
    Parser de documentos Word (.docx) con extracción de:
    - Texto completo con formato preservado
    - Tablas
    - Metadatos del documento
    - Estilos y encabezados (para chunking por sección)
    """

    def __init__(self):
        self._available = False
        try:
            import docx
            self._docx = docx
            self._available = True
        except ImportError:
            pass

    async def extract_text(self, file_path: str) -> str:
        """Extrae todo el texto del documento."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
        if not self._available:
            raise ImportError("python-docx no está instalado. pip install python-docx")

        doc = self._docx.Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)

        return "\n".join(text_parts)

    async def extract_structured(self, file_path: str) -> Dict[str, Any]:
        """
        Extrae el documento de forma estructurada, preservando:
        - Encabezados (Heading 1, 2, 3)
        - Párrafos normales
        - Tablas
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
        if not self._available:
            raise ImportError("python-docx no está instalado")

        doc = self._docx.Document(file_path)
        sections = []
        current_section = {"heading": None, "content": [], "level": 0}

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name if paragraph.style else "Normal"

            if style_name.startswith("Heading"):
                # Guardar sección anterior
                if current_section["content"]:
                    sections.append(current_section)

                level = int(style_name.replace("Heading ", "")) if "Heading" in style_name else 1
                current_section = {
                    "heading": text,
                    "content": [],
                    "level": level,
                }
            else:
                current_section["content"].append({
                    "text": text,
                    "style": style_name,
                    "bold": paragraph.runs[0].bold if paragraph.runs else False,
                })

        # Última sección
        if current_section["content"]:
            sections.append(current_section)

        # Extraer tablas
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        return {
            "paragraphs": len(doc.paragraphs),
            "sections": sections,
            "tables": tables,
            "metadata": await self.get_metadata(file_path),
        }

    async def extract_tables(self, file_path: str) -> List[List[List[str]]]:
        """
        Extrae todas las tablas del documento.
        Returns: Lista de tablas, cada tabla es lista de filas, cada fila es lista de celdas.
        """
        if not self._available:
            return []

        doc = self._docx.Document(file_path)
        tables = []

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        return tables

    async def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Obtiene metadatos del documento Word.
        """
        if not self._available:
            return {"filename": os.path.basename(file_path)}

        doc = self._docx.Document(file_path)
        props = doc.core_properties

        return {
            "filename": os.path.basename(file_path),
            "file_size": os.path.getsize(file_path),
            "author": props.author,
            "title": props.title,
            "subject": props.subject,
            "created": str(props.created) if props.created else None,
            "modified": str(props.modified) if props.modified else None,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }
