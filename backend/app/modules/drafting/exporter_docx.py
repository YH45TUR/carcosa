"""
Sistema Legal CO - Exportador DOCX
Exporta documentos legales a formato Word (.docx) con norma NTC 1486.

Formato NTC 1486 (Norma Técnica Colombiana para documentación):
- Fuente: Times New Roman 12pt
- Márgenes: Superior 3cm, Izquierda 4cm, Derecha 2cm, Inferior 3cm
- Interlineado: 1.5 líneas
- Numeración legal de páginas
- Firma con T.P. (Tarjeta Profesional)
"""
from typing import Optional, Dict, Any
import os


class DOCXExporter:
    """
    Exportador de documentos legales a formato .docx siguiendo NTC 1486.
    """

    def __init__(self):
        self._available = False
        try:
            import docx
            from docx.shared import Pt, Cm, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_ORIENT
            self._docx = docx
            self._Pt = Pt
            self._Cm = Cm
            self._Inches = Inches
            self._RGBColor = RGBColor
            self._WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
            self._WD_ORIENT = WD_ORIENT
            self._available = True
        except ImportError:
            pass

    async def export(
        self,
        content: str,
        output_path: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        Exporta contenido de texto plano a un archivo .docx con formato NTC 1486.

        Args:
            content: Texto del documento legal
            output_path: Ruta donde guardar el .docx
            metadata: Metadatos (título, autor, etc.)

        Returns:
            Ruta al archivo generado
        """
        if not self._available:
            raise ImportError(
                "python-docx no está instalado. pip install python-docx"
            )

        doc = self._docx.Document()

        # Configurar márgenes NTC 1486
        section = doc.sections[0]
        section.top_margin = self._Cm(3)
        section.bottom_margin = self._Cm(3)
        section.left_margin = self._Cm(4)
        section.right_margin = self._Cm(2)

        # Estilo base Times New Roman 12pt, interlineado 1.5
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = self._Pt(12)
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_after = self._Pt(0)
        paragraph_format.space_before = self._Pt(0)

        # Procesar el contenido línea por línea
        lines = content.split('\n')
        in_list = False

        for line in lines:
            line_stripped = line.strip()

            if not line_stripped:
                # Línea vacía = salto de párrafo
                doc.add_paragraph('')
                in_list = False
                continue

            # Detectar encabezados (SECCIONES EN MAYÚSCULAS)
            if line_stripped.isupper() and len(line_stripped) > 3:
                p = doc.add_paragraph()
                run = p.add_run(line_stripped)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = self._Pt(12)
                p.alignment = self._WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = self._Pt(12)
                p.paragraph_format.space_after = self._Pt(6)
                in_list = False
                continue

            # Detectar numeración legal (I., II., 1., 2., a), b), etc.)
            if (line_stripped[0].isdigit() and '. ' in line_stripped[:5]) or \
               (len(line_stripped) > 2 and line_stripped[0].isalpha() and ') ' in line_stripped[:5]):
                p = doc.add_paragraph()
                run = p.add_run(line_stripped)
                run.font.name = 'Times New Roman'
                run.font.size = self._Pt(12)
                p.paragraph_format.left_indent = self._Cm(1)
                p.paragraph_format.first_line_indent = self._Cm(-0.5)
                in_list = True
                continue

            # Texto normal
            p = doc.add_paragraph()
            run = p.add_run(line_stripped)
            run.font.name = 'Times New Roman'
            run.font.size = self._Pt(12)
            p.paragraph_format.first_line_indent = self._Cm(1)
            in_list = False

        # Agregar metadatos
        if metadata:
            doc.core_properties.title = metadata.get('title', '')
            doc.core_properties.author = metadata.get('author', '')
            doc.core_properties.subject = metadata.get('subject', '')

        # Guardar
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        doc.save(output_path)

        return output_path

    async def export_with_template(
        self,
        content: str,
        output_path: str,
        document_type: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        Exporta con formato específico según el tipo de documento legal.

        Args:
            content: Texto del documento
            output_path: Ruta de salida
            document_type: Tipo (demanda, tutela, recurso, etc.)
            metadata: Metadatos adicionales

        Returns:
            Ruta al archivo generado
        """
        doc = self._docx.Document()

        section = doc.sections[0]
        section.top_margin = self._Cm(3)
        section.bottom_margin = self._Cm(3)
        section.left_margin = self._Cm(4)
        section.right_margin = self._Cm(2)

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = self._Pt(12)
        style.paragraph_format.line_spacing = 1.5

        # Agregar encabezado con membrete legal si es necesario
        if document_type in ('demanda', 'tutela'):
            header = section.header
            hp = header.paragraphs[0]
            hp.text = 'SEÑOR JUEZ'
            hp.alignment = self._WD_ALIGN_PARAGRAPH.RIGHT

        # Procesar contenido sección por sección
        lines = content.split('\n')
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                doc.add_paragraph('')
                continue

            p = doc.add_paragraph()
            run = p.add_run(line_stripped)

            # Detectar secciones en mayúsculas sostenidas
            if line_stripped.isupper() and len(line_stripped) > 3:
                run.bold = True
                p.alignment = self._WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = self._Pt(12)
            else:
                p.paragraph_format.first_line_indent = self._Cm(1)

            run.font.name = 'Times New Roman'
            run.font.size = self._Pt(12)

        # Footer con numeración automática de páginas
        from docx.oxml.ns import qn
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = self._WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run('Página ')
        run.font.name = 'Times New Roman'
        run.font.size = self._Pt(10)
        # Agregar campo de número de página automático
        fldChar1 = self._docx.oxml.OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        instrText = self._docx.oxml.OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run._r.append(instrText)
        fldChar2 = self._docx.oxml.OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

        # Agregar metadatos
        if metadata:
            doc.core_properties.title = metadata.get('title', document_type)
            doc.core_properties.author = metadata.get('author', '')
            if metadata.get('comments'):
                doc.core_properties.comments = metadata['comments']

        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        doc.save(output_path)

        return output_path
